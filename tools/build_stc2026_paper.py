from __future__ import annotations

import json
import sys
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MODEL_META = ROOT / "models" / "AutogluonModels" / "TCN_aug_weighted_v1" / "model_meta.json"
METRICS = ROOT / "models" / "AutogluonModels" / "TCN_aug_weighted_v1" / "metrics.json"
FEATURE_SUMMARY = ROOT / "analysis" / "feature_importance_TCN_aug_weighted_v1" / "feature_summary.txt"


def set_font(run, name="Arial Narrow", size=12, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="9DB5C8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_cell_margins(table, margin_dxa=90):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side in ["top", "left", "bottom", "right"]:
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(margin_dxa))
        node.set(qn("w:type"), "dxa")


def clear_document(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_heading(doc: Document, text: str, before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_font(r, size=12, bold=True, color=(0, 58, 112))
    return p


def add_body(doc: Document, text: str, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(text)
    set_font(r, size=12)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, size=10, italic=True)
    return p


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    r = p.add_run(text)
    set_font(r, size=12)
    return p


def fill_paragraph(cell, text: str, bold=False, size=12, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.clear()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def math_run(text: str, style="i") -> str:
    return (
        f'<m:r><m:rPr><m:sty m:val="{style}"/></m:rPr>'
        f"<m:t>{escape(text)}</m:t></m:r>"
    )


def math_sub(base: str, sub: str, base_style="i", sub_style="i") -> str:
    return (
        "<m:sSub>"
        f"<m:e>{math_run(base, base_style)}</m:e>"
        f"<m:sub>{math_run(sub, sub_style)}</m:sub>"
        "</m:sSub>"
    )


def math_sup(base: str, sup: str, base_style="i", sup_style="p") -> str:
    return (
        "<m:sSup>"
        f"<m:e>{math_run(base, base_style)}</m:e>"
        f"<m:sup>{math_run(sup, sup_style)}</m:sup>"
        "</m:sSup>"
    )


def math_subsup(base: str, sub: str, sup: str, base_style="p", sub_style="i", sup_style="i") -> str:
    return (
        "<m:sSubSup>"
        f"<m:e>{math_run(base, base_style)}</m:e>"
        f"<m:sub>{math_run(sub, sub_style)}</m:sub>"
        f"<m:sup>{math_run(sup, sup_style)}</m:sup>"
        "</m:sSubSup>"
    )


def math_frac(num: str, den: str) -> str:
    return f"<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>"


def add_omml_equation(doc: Document, inner_omml: str, number: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    xml = (
        f'<m:oMathPara {nsdecls("m")}>'
        "<m:oMathParaPr><m:jc m:val=\"centerGroup\"/></m:oMathParaPr>"
        "<m:oMath>"
        f"{inner_omml}"
        f'{math_run("        " + number, "p")}'
        "</m:oMath>"
        "</m:oMathPara>"
    )
    p._p.append(parse_xml(xml))
    return p


def add_title_block(doc: Document):
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    set_table_width(table, 9360)
    set_table_borders(table, color="D4DDE9", size="4")
    set_table_cell_margins(table, margin_dxa=95)
    widths = [1650, 7710]
    rows = [
        ("TITLE", "Adaptive Wait-Time Prediction for QOET TTR Waveforms Using a TCN-AutoGluon Hybrid Model"),
        ("AUTHORS", "Teerawit Pongkunawut and Sukit Saelao"),
        ("AREA", "Software & Algorithm; Test and Measurement; Quality and Reliability"),
        ("KEYWORDS", "QOET, TTR, waveform regression, Temporal Convolutional Network, AutoGluon, MLflow"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, widths[idx])
        set_cell_shading(row.cells[0], "E8F6FD")
        fill_paragraph(row.cells[0], label, bold=True, size=10, color=(0, 58, 112))
        fill_paragraph(row.cells[1], value, bold=(label == "TITLE"), size=12)
    doc.add_paragraph()


def add_flow_table(doc: Document):
    table = doc.add_table(rows=2, cols=4)
    table.autofit = False
    set_table_width(table, 9360)
    set_table_borders(table, color="8FA7BC", size="5")
    set_table_cell_margins(table, margin_dxa=85)
    widths = [2160, 2400, 2400, 2400]
    headers = ["1. Data", "2. Features", "3. Learning", "4. Review"]
    values = [
        "SignalSample CSV/XLSX converted to long waveform rows",
        "Robust waveform features plus fixed-length tensors",
        "TCN embeddings merged with AutoGluon tabular ensemble",
        "MLflow registry, diagnostics, web prediction dashboard",
    ]
    for c, header in enumerate(headers):
        set_cell_width(table.rows[0].cells[c], widths[c])
        set_cell_shading(table.rows[0].cells[c], "003A70")
        fill_paragraph(table.rows[0].cells[c], header, bold=True, size=10, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_width(table.rows[1].cells[c], widths[c])
        fill_paragraph(table.rows[1].cells[c], values[c], size=10)
    add_caption(doc, "Figure 1. Simplified process flow for the QOET TTR wait-time prediction pipeline.")


def add_results_table(doc: Document, meta: dict, metrics: dict):
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    set_table_width(table, 9360)
    set_table_borders(table)
    set_table_cell_margins(table, margin_dxa=85)
    widths = [1850, 2100, 2100, 3310]
    headers = ["Metric", "Validation", "Test", "Engineering meaning"]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "003A70")
        fill_paragraph(cell, headers[idx], bold=True, size=10, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

    m = meta["result"]["metrics"]
    valid = metrics.get("valid", {})
    test = metrics.get("test", {})
    rows = [
        ("MAE", f"{valid.get('mae', 0):.6f} ms", f"{test.get('mae', m['mae_all']):.6f} ms", "Average absolute timing error on held-out waveforms."),
        ("RMSE", f"{valid.get('rmse', 0):.6f} ms", f"{test.get('rmse', m['rmse']):.6f} ms", "Higher penalty for large misses; useful for tail-risk review."),
        ("Fast precision", "-", f"{m['fast_precision']:.3f}", "No false fast calls at the 0.1 ms threshold in this run."),
        ("Fast recall", "-", f"{m['fast_recall']:.3f}", "Half of true fast events were detected; the main tuning target."),
        ("Best TCN epoch", "-", str(meta["result"]["overfitting_summary"]["best_epoch"]), "Train and validation loss converged with Good Fit status."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_width(cells[idx], widths[idx])
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in [1, 2] else WD_ALIGN_PARAGRAPH.LEFT
            fill_paragraph(cells[idx], value, bold=(idx == 0), size=10, align=align)
    add_caption(doc, "Table 1. Validation and test performance for TCN_aug_weighted_v1.")


def build_doc(template_path: Path, out_path: Path):
    meta = json.loads(MODEL_META.read_text(encoding="utf-8"))
    metrics = json.loads(METRICS.read_text(encoding="utf-8"))

    doc = Document(template_path)
    clear_document(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Narrow"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Narrow")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Narrow")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial Narrow")
    normal.font.size = Pt(12)

    add_title_block(doc)

    add_heading(doc, "Abstract", before=4)
    add_body(
        doc,
        "This paper presents a production-oriented workflow for predicting wait_time_ms from QOET TTR waveform signals. The solution converts raw SignalSample-style wide files into long waveform records, extracts robust waveform-shape features, learns temporal embeddings using a Temporal Convolutional Network (TCN), and trains an AutoGluon tabular ensemble on the combined feature space. The current candidate model, TCN_aug_weighted_v1, achieved test MAE = 0.165512 ms and RMSE = 0.289210 ms on 1,000 held-out waveforms, with fast-event precision of 1.000 at the 0.1 ms threshold. The key contribution is a repeatable hybrid pipeline that combines signal-processing transparency with learned waveform representation, giving QOET engineers a practical path from raw measurement files to prediction, diagnostics, model registry, and waveform-level review.",
    )

    add_heading(doc, "Background")
    add_body(
        doc,
        "QOET TTR analysis depends on interpreting waveform settling behavior and estimating the wait time at which a signal reaches a useful operating state. Manual review is slow, inconsistent, and difficult to scale when the data set contains thousands of waveforms with different amplitudes, noise floors, late transitions, rebounds, and fast-zero cases. The project supports software & algorithm, test and measurement, and quality/reliability goals by converting repeated waveform review into a measurable, traceable ML pipeline.",
    )
    add_body(
        doc,
        "The technical challenge was not only regression accuracy. The model also needed to preserve explainability for engineering review, support raw CSV/XLSX files collected from instruments, separate train/validation/test splits, track artifacts in MLflow, and produce diagnostics that make failure modes visible. A pure black-box model would be difficult to trust, while only handcrafted features could miss waveform morphology that is easy to see but hard to encode. This motivated a hybrid design: engineered features capture known settling behavior while TCN embeddings learn temporal patterns directly from the waveform.",
    )

    add_page_break(doc)
    add_heading(doc, "Project Description", before=0)
    add_body(
        doc,
        "The implemented workflow has four stages. First, the application accepts raw SignalSample CSV/XLSX files and converts wide instrument columns into a long waveform table. Second, a feature pipeline computes robust timing, slope, quiet-window, tail-creep, ringing, overshoot, plateau, and late-transition features. Third, a TCN is trained on fixed-length waveform tensors and exports a learned embedding for each waveform. Finally, AutoGluon trains a tabular ensemble on the merged engineered and learned feature space, while MLflow records metrics, artifacts, registry status, and model alias information.",
    )
    add_flow_table(doc)
    add_body(
        doc,
        "The TCN design follows the forecasting idea described by Unit8: a TCN uses causal one-dimensional convolution so an output time step only depends on current and previous inputs, and dilation expands the receptive field without requiring a very deep recurrent model. This is appropriate for TTR waveforms because settling behavior is local in time but can also depend on earlier shape, amplitude, and late-window history.",
    )
    add_body(doc, "The raw-to-feature conversion is expressed below. The equations are stored as Word Office Math objects, not ordinary text runs.", after=2)

    add_omml_equation(
        doc,
        math_sub("D", "long")
        + math_run(" = {(", "p")
        + math_sub("wave", "id", "p", "p")
        + math_run(" = k, sample = i, ", "p")
        + math_sub("t", "ms")
        + math_run(" = iΔt, value = ", "p")
        + math_sub("x", "k,i")
        + math_run(")}", "p"),
        "(1)",
    )
    add_omml_equation(
        doc,
        math_sub("x̃", "k,i")
        + math_run(" = clip((", "p")
        + math_sub("x", "k,i")
        + math_run(" - ", "p")
        + math_sub("r", "end,k")
        + math_run(") / ", "p")
        + math_sub("d", "k")
        + math_run(", -12, 12)", "p"),
        "(2)",
    )
    add_omml_equation(
        doc,
        math_sub("d", "k")
        + math_run(" = max(|", "p")
        + math_sub("r", "end,k")
        + math_run(" - ", "p")
        + math_sub("r", "start,k")
        + math_run("|, 0.12", "p")
        + math_sub("s", "k")
        + math_run(", 6", "p")
        + math_sub("σ", "end,k")
        + math_run(", ε)", "p"),
        "(3)",
    )
    add_omml_equation(
        doc,
        math_sub("H", "k")
        + math_run(" = ", "p")
        + math_sub("TCN", "θ", "p", "i")
        + math_run("(", "p")
        + math_sub("X", "k")
        + math_run("),  ", "p")
        + math_sub("z", "k")
        + math_run(" = ReLU(", "p")
        + math_sub("W", "e")
        + math_run("GAP(", "p")
        + math_sub("H", "k")
        + math_run(") + ", "p")
        + math_sub("b", "e")
        + math_run("),  ", "p")
        + math_sub("ŷ", "TCN,k")
        + math_run(" = ", "p")
        + math_sub("W", "h")
        + math_sub("z", "k")
        + math_run(" + ", "p")
        + math_sub("b", "h"),
        "(4)",
    )
    add_omml_equation(
        doc,
        math_sub("L", "TCN")
        + math_run(" = ", "p")
        + math_frac(math_run("1", "p"), math_run("N", "i"))
        + math_subsup("∑", "k=1", "N", "p", "p", "i")
        + math_sub("w", "k")
        + math_run("(", "p")
        + math_sub("ŷ", "log,k")
        + math_run(" - log(1 + ", "p")
        + math_sub("y", "k")
        + math_run("))", "p")
        + math_run("²", "p"),
        "(5)",
    )
    add_omml_equation(
        doc,
        math_sub("F", "k")
        + math_run(" = [", "p")
        + math_sub("φ", "k")
        + math_run("; ", "p")
        + math_sub("z", "k")
        + math_run("],  ", "p")
        + math_sub("ŷ", "k")
        + math_run(" = exp(", "p")
        + math_sub("f", "AG")
        + math_run("(", "p")
        + math_sub("F", "k")
        + math_run(")) - 1", "p"),
        "(6)",
    )
    add_body(
        doc,
        "Equations (1)-(3) describe input conversion and robust normalization; equations (4)-(6) describe the learned TCN representation, the weighted log-target training objective, and the final AutoGluon regression target. The web application exposes this same workflow for upload, training, reuse of existing TCN models, model registry inspection, and direct prediction from new SignalSample files.",
    )

    add_page_break(doc)
    add_heading(doc, "Results, Learnings or Take-Aways", before=0)
    add_results_table(doc, meta, metrics)
    feature = meta["result"]["feature_summary"]
    add_body(
        doc,
        f"The feature-importance audit supports the hybrid design. Out of {feature['total_features']} total features, the top-{feature['topn']} set contained {feature['top30_count']['tcn_embedding']} TCN embedding dimensions, {feature['top30_count']['late_settle']} late-settle features, and {feature['top30_count']['handcrafted_other']} other handcrafted features. Summed importance was {feature['group_sum']['tcn_embedding']:.6f} for TCN embeddings, {feature['group_sum']['late_settle']:.6f} for late-settle features, and {feature['group_sum']['handcrafted_other']:.6f} for other handcrafted features. The top-10 list was dominated by learned embeddings, while num_mean_shifts also appeared, confirming that delayed-transition behavior remains useful as an engineered signal.",
    )
    add_body(
        doc,
        "The main learning is that waveform timing prediction benefits from both learned representation and explicit settling features. Pure handcrafted features are easier to explain but can miss subtle waveform morphology. A TCN embedding alone is stronger but less auditable. The hybrid model gives a useful compromise: accuracy from learned waveform representation, engineering trust from interpretable late-settle features, and operational control from MLflow artifacts, model registry metadata, and waveform plots.",
    )
    add_body(
        doc,
        "The current limitation is fast-event recall. Precision at the 0.1 ms threshold is perfect in this run, but recall is 0.500, meaning the system is conservative about declaring a fast event. That behavior is acceptable when avoiding false fast calls is more important than catching every fast waveform, but future work should explore class-aware thresholding, stronger fast-event weighting, and calibration of the final prediction distribution.",
    )
    add_body(
        doc,
        "A practical takeaway for reuse is that the data contract matters as much as the model. Once raw files are reliably converted into a stable long-waveform representation, the same code path can train, validate, backtest, explain feature importance, and run inference. This reduces handoff risk between notebook experimentation and production-style engineering review.",
    )

    add_page_break(doc)
    add_heading(doc, "Conclusion", before=0)
    add_bullet(doc, "A complete QOET TTR ML workflow was implemented from raw SignalSample input to model training, prediction, diagnostics, and registry tracking.")
    add_bullet(doc, "The TCN-AutoGluon hybrid candidate reached test MAE = 0.165512 ms and RMSE = 0.289210 ms on 1,000 held-out waveforms, with fast precision = 1.000.")
    add_bullet(doc, "Feature importance shows the model uses both learned TCN embeddings and late-settle engineered features, validating the hybrid approach.")
    add_body(
        doc,
        "Future work will focus on improving fast-event recall, adding calibrated confidence bands, expanding tests on real production lots, and packaging the prediction workflow behind role-based access so training operations are restricted to authorized engineers.",
    )

    add_heading(doc, "Acknowledgements")
    add_body(
        doc,
        "The authors acknowledge the QOET engineering context, project reviewers, and contributors who supported data preparation, model validation, dashboard review, and workflow feedback. Development credit: Teerawit Pongkunawut and Sukit Saelao.",
    )

    add_heading(doc, "References")
    refs = [
        "F. Lässig, \"Temporal Convolutional Networks and Forecasting,\" Unit8, 2021. https://unit8.com/resources/temporal-convolutional-networks-and-forecasting/",
        "S. Bai, J. Z. Kolter, and V. Koltun, \"An Empirical Evaluation of Generic Convolutional and Recurrent Networks for Sequence Modeling,\" arXiv:1803.01271, 2018.",
        "Project artifacts: scripts/data, scripts/features, scripts/tcn, scripts/autogluon, MLflow run 85eac2bff0d048e89bc98987896f3420.",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    template = Path(sys.argv[1])
    out = Path(sys.argv[2])
    build_doc(template, out)


if __name__ == "__main__":
    main()
