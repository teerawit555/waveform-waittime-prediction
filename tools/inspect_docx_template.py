from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def main() -> None:
    path = Path(sys.argv[1])
    doc = Document(path)

    print(f"DOCX: {path}")
    print(f"Sections: {len(doc.sections)}")
    for i, section in enumerate(doc.sections, start=1):
        print(
            "SECTION",
            i,
            "page",
            section.page_width,
            section.page_height,
            "margins",
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
        )

    print("\nPARAGRAPHS")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"[{i:03d}] style={p.style.name!r}: {text}")

    print("\nTABLES")
    for ti, table in enumerate(doc.tables):
        print(f"TABLE {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri, row in enumerate(table.rows[:10]):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            print(f"  r{ri}: {cells}")


if __name__ == "__main__":
    main()
